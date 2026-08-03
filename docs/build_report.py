"""
Convert docs/PROJECT_REPORT.md -> Foodly_COS30043_Project_Report.docx
Two-pass: parse markdown into an element list, then render
- title page
- static Table of Contents
- body (headings, figures w/ captions, tables, code blocks, lists, paragraphs)
"""
import os
import re

ROOT = r"C:\Users\vieth\COS30043 final project"
SRC = os.path.join(ROOT, "docs", "PROJECT_REPORT.md")
FIG_DIR = os.path.join(ROOT, "docs", "figures")
# Write to a fresh filename to avoid OneDrive locks on the live docx:
OUT = os.path.join(ROOT, "Foodly_COS30043_Project_Report.docx")
FALLBACK_OUT = os.path.join(ROOT, "Foodly_COS30043_Project_Report_GENERATED.docx")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for lvl, sz in [("Heading 1", Pt(22)), ("Heading 2", Pt(16)), ("Heading 3", Pt(14)),
                ("Heading 4", Pt(12)), ("Heading 5", Pt(11)), ("Heading 6", Pt(10))]:
    h = doc.styles[lvl]
    h.font.name = "Calibri"
    h.font.size = sz
    h.font.bold = True

if "CodeBlock" not in [s.name for s in doc.styles]:
    CODE = doc.styles.add_style("CodeBlock", 1)
else:
    CODE = doc.styles["CodeBlock"]
CODE.font.name = "Consolas"
CODE.font.size = Pt(9)
CODE.paragraph_format.before = Pt(6)
CODE.paragraph_format.after = Pt(6)
CODE.paragraph_format.left_indent = Pt(24)
CODE.paragraph_format.right_indent = Pt(24)
CODE.paragraph_format.line_spacing = 1.0

CAP = doc.styles["Caption"]
CAP.font.name = "Calibri"; CAP.font.size = Pt(9); CAP.font.italic = True
CAP.paragraph_format.space_before = Pt(2); CAP.paragraph_format.space_after = Pt(6)
CAP.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -------------------------------- parser ----------------------------------
class F:  # Figure element
    def __init__(self, src, caption):
        self.kind = "figure"
        self.src = src
        self.caption = caption
    def render(self, doc):
        full = self.src
        if not os.path.isabs(full):
            full = os.path.join(FIG_DIR, full)
        if not os.path.isfile(full):
            p = doc.add_paragraph("[Figure not found: {}]".format(self.src))
            p.runs[0].italic = True
            return
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(full, width=Inches(6.2))
        if self.caption:
            doc.add_paragraph(self.caption, style="Caption")
class H:
    def __init__(self, level, text):
        self.kind = "heading"; self.level = level; self.text = text
    def render(self, doc):
        doc.add_heading(self.text, level=min(self.level, 3))
class T:
    def __init__(self, rows):
        self.kind = "table"
        self.rows = rows
    def render(self, doc):
        if not self.rows: return
        t = doc.add_table(rows=len(self.rows), cols=max(len(r) for r in self.rows))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(self.rows):
            for ci in range(len(self.rows[0])):
                t.rows[ri].cells[ci].text = row[ci].strip() if ci < len(row) else ""
class C:
    def __init__(self, lines):
        self.kind = "code"; self.lines = lines
    def render(self, doc):
        doc.add_paragraph("\n".join(self.lines), style="CodeBlock")
class L:
    def __init__(self, kind, text):
        self.kind = "list"; self.listkind = kind; self.text = text  # 'bullet' or 'number'
    def render(self, doc):
        p = doc.add_paragraph(style="List Bullet" if self.listkind == "bullet" else "List Number")
        add_runs(p, self.text)
class P:
    def __init__(self, text):
        self.kind = "para"; self.text = text
    def render(self, doc):
        p = doc.add_paragraph()
        add_runs(p, self.text)

def add_runs(par, text):
    tokens = re.split(r"(?<!`)`([^`]+)`(?!`)", text)
    for ti, tok in enumerate(tokens):
        if ti % 2 == 1:
            r = par.add_run(tok); r.font.name = "Consolas"; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37); continue
        parts = re.split(r"(\*\*[^*]+\*\*)", tok)
        for p in parts:
            if not p: continue
            if p.startswith("**") and p.endswith("**"):
                r = par.add_run(p[2:-2]); r.bold = True
            else:
                mm = re.match(r"_(.+?)_", p)
                if mm:
                    r = par.add_run(mm.group(1)); r.italic = True
                else:
                    par.add_run(p)

# parse
raw = open(SRC, encoding="utf-8").read().splitlines()
title_line = raw[0][2:].strip() if raw else "Foodly — Project Report"
# collect meta (subtitle lines) until a '---' rule
meta, i = [], 1
while i < len(raw) and raw[i].strip() and not re.match(r"^-{3,}$", raw[i]):
    meta.append(raw[i].strip()); i += 1
# skip rules/blanks
while i < len(raw) and (raw[i].strip() == "" or re.match(r"^-{3,}$", raw[i])):
    i += 1

elements = []
TOC = []
in_code = False
code_buf = []
in_table = False
tab_buf = []
pending_toc = None  # (level, text)

def flush_table():
    global in_table, tab_buf
    if in_table and tab_buf:
        elements.append(T(tab_buf))
    tab_buf = []
    in_table = False

def flush_code():
    global in_code, code_buf
    if in_code:
        elements.append(C(code_buf))
    code_buf = []
    in_code = False

# token types: ('text', s) ('fig', fn, caption)
FIG_RE = re.compile(r"(!\[\[(?P<fn>[^\]]+)\])|(!\[(?P<cap>[^\]]*)\]\((?P<src>[^)]+)\))")
def split_inline(line):
    out = []
    pos = 0
    for m in FIG_RE.finditer(line):
        if m.start() > pos:
            out.append(P(line[pos:m.start()].strip()))
        if m.group("fn") is not None:  # ![[fn]]
            fn = m.group("fn")
            out.append(F(fn, os.path.basename(fn)))
        elif m.group("src") is not None:  # ![cap](path)
            out.append(F(m.group("src"), m.group("cap") or ""))
        pos = m.end()
    if pos < len(line):
        rest = line[pos:].strip()
        if rest:
            out.append(P(rest))
    return out

while i < len(raw):
    line = raw[i]
    if re.match(r"^\s*```", line):
        if in_code:
            elements.append(C(code_buf)); code_buf = []; in_code = False
        else:
            flush_table(); in_code = True; code_buf = []
        i += 1; continue
    if in_code:
        code_buf.append(line); i += 1; continue

    if line.strip().startswith("|"):
        cells = [re.sub(r"`?\!?\[\[([^\]]+)\]\]?", r"\1", c.strip())
                 for c in re.split(r"\|", line.strip().strip("|")) if c != ""]
        if not cells:
            i += 1; continue
        if re.match(r"^[\s:|-]+$", " ".join(cells)):
            i += 1; continue
        if not in_table:
            in_table = True; tab_buf = [cells]
        else:
            tab_buf.append(cells)
        i += 1; continue

    # flush table when next line is not a table row
    if in_table and not line.strip().startswith("|") and not re.match(r"^-{3,}$", line):
        flush_table()

    if re.match(r"^-{3,}\s*$", line):
        i += 1; continue

    if line.strip() == "":
        i += 1; continue

    if line.startswith("#"):
        flush_table()
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            lvl = len(m.group(1)); text = m.group(2)
            elements.append(H(lvl, text))
            TOC.append((lvl, text))
        i += 1; continue

    if re.match(r"^\s*[-*] ", line):
        elements.append(L("bullet", re.sub(r"^\s*[-*] ", "", line)))
        i += 1; continue
    if re.match(r"^\s*\d+\.\s", line):
        elements.append(L("number", re.sub(r"^\s*\d+\.\s*", "", line)))
        i += 1; continue

    flush_table()
    for el in split_inline(line):
        elements.append(el)
    i += 1

flush_table(); flush_code()

# ---- render: title page ----
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(title_line)
r.font.name = "Calibri"; r.font.size = Pt(26); r.bold = True
for _ in range(4):
    doc.add_paragraph()
for m in meta:
    if m.startswith("---"): continue
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(m)
    r.font.name = "Calibri"; r.font.size = Pt(13)
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("COS30043 — Interface Design and Development\nSwinburne University of Technology")
r.font.name = "Calibri"; r.font.size = Pt(12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Submitted: August 2026")
r.font.name = "Calibri"; r.font.size = Pt(11)
doc.add_page_break()

# ---- render: Table of Contents ----
doc.add_paragraph("Table of Contents", style="Heading 1")
for lvl, text in TOC:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12 * (lvl - 1))
    # dots leader is hard in python-docx without fields; use a simple leader
    label = "." * max(1, 50 - len(text) - (lvl - 1) * 3)
    r = p.add_run("{} {}".format(text, label))
    r.font.name = "Calibri"; r.font.size = Pt(11)
    # page number placeholder: use a RIGHT-aligned tab
    p.add_run("    ...").font.name = "Calibri"
doc.add_page_break()

# ---- render body ----
for e in elements:
    e.render(doc)

target = OUT
try:
    doc.save(target)
except (PermissionError, OSError):
    # OneDrive lock on the live docx: write to a fresh name instead
    target = FALLBACK_OUT
    doc.save(target)
    print("NOTE: Original docx was locked by OneDrive; written to a fresh file instead.")
print("Saved:", target)
print("Elements:", len(elements), "TOC:", len(TOC), "Headings in TOC:", len(TOC))
